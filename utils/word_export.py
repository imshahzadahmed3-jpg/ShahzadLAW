import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, fill, text_color=None):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    if text_color:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(text_color)

def export_to_word(df, totals, filename="transactions_report.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading('Bank Calculation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary Section
    doc.add_heading('Grand Totals', level=1)
    
    # Create a nice summary table
    summary_table = doc.add_table(rows=2, cols=3)
    summary_table.style = 'Table Grid'
    
    headers = ['Total Received (In)', 'Total Sent (Out)', 'Tax / Charges']
    vals = [f"Rs. {totals['credit']:,.2f}", f"Rs. {totals['debit']:,.2f}", f"Rs. {totals['tax']:,.2f}"]
    
    for i, h in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'E2E8F0')
        
    for i, v in enumerate(vals):
        cell = summary_table.rows[1].cells[i]
        cell.text = v
        cell.paragraphs[0].runs[0].bold = True
        if i == 0:
            set_cell_background(cell, 'DCFCE7') # Greenish
        elif i == 1:
            set_cell_background(cell, 'FEE2E2') # Reddish
        else:
            set_cell_background(cell, 'FEF3C7') # Yellowish

    doc.add_paragraph()
    
    # Bank Breakdown
    doc.add_heading('Bank Breakdown', level=2)
    bd_table = doc.add_table(rows=3, cols=3)
    bd_table.style = 'Table Grid'
    
    bd_headers = ['Bank', 'Received (In)', 'Sent (Out)']
    for i, h in enumerate(bd_headers):
        cell = bd_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'E2E8F0')
        
    row1 = bd_table.rows[1].cells
    row1[0].text = 'Meezan Bank'
    row1[1].text = f"Rs. {totals.get('meezan_credit', 0):,.2f}"
    row1[2].text = f"Rs. {totals.get('meezan_debit', 0):,.2f}"
    
    row2 = bd_table.rows[2].cells
    row2[0].text = 'Alfalah Bank'
    row2[1].text = f"Rs. {totals.get('alfalah_credit', 0):,.2f}"
    row2[2].text = f"Rs. {totals.get('alfalah_debit', 0):,.2f}"
    
    doc.add_paragraph()
    
    # Table Section
    doc.add_heading('Transactions Details', level=1)
    
    if df.empty:
        doc.add_paragraph("No transactions found.")
    else:
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        headers = ['Date', 'Bank', 'Account', 'Party Name', 'Description', 'Debit (Out)', 'Credit (In)', 'Source File']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background(hdr_cells[i], '1E293B', 'FFFFFF')
            
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['transaction_date'])
            row_cells[1].text = str(row['bank'])
            row_cells[2].text = str(row.get('account_number', ''))
            row_cells[3].text = str(row['party_name'])
            row_cells[4].text = str(row['description'])[:60] # Truncate long descriptions
            
            debit_val = row['debit']
            row_cells[5].text = f"{debit_val:,.2f}"
            if debit_val > 0:
                set_cell_background(row_cells[5], 'FEE2E2')
                
            credit_val = row['credit']
            row_cells[6].text = f"{credit_val:,.2f}"
            if credit_val > 0:
                set_cell_background(row_cells[6], 'DCFCE7')

            row_cells[7].text = str(row.get('source_file', 'Manual'))
            
    doc.save(filename)
    return filename
