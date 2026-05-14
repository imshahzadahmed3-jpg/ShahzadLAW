import docx

doc = docx.Document("Owais_Accounts_With_Transaction_IDs.docx")

print("--- TEXT ---")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)

print("\n--- TABLES ---")
for i, table in enumerate(doc.tables):
    print(f"Table {i+1}:")
    for j, row in enumerate(table.rows):
        row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        print(row_data)
        if j > 10:
            print("...")
            break
